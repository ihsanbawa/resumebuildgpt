#!/usr/bin/env python
"""
resume_build.py  –  Replace placeholders in a Word template, output DOCX + PDF.

Usage (local test):
    python resume_build.py \
        --template template.docx \
        --json     data.json \
        --out      /tmp/resume_out

Produces:
    /tmp/resume_out.docx      (always)
    /tmp/resume_out.pdf       (if docx2pdf or LibreOffice is available)
"""

import argparse, json, sys, subprocess
from pathlib import Path
from docx import Document

# ---------- helper functions ----------
def get_all_paragraphs(doc):
    paragraphs = []
    def iter_paragraphs(parent):
        for p in parent.paragraphs:
            paragraphs.append(p)
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    iter_paragraphs(cell)
    iter_paragraphs(doc)
    return paragraphs

def replace_text_keeping_style(paragraph, placeholder, replacement):
    runs = paragraph.runs
    full_text = ''.join(run.text for run in runs)
    idx = full_text.find(placeholder)
    while idx != -1:
        cum = 0
        s_run = s_idx = e_run = e_idx = None
        for i, run in enumerate(runs):
            r_len = len(run.text)
            if s_run is None and cum + r_len > idx:
                s_run, s_idx = i, idx - cum
            if e_run is None and cum + r_len >= idx + len(placeholder):
                e_run, e_idx = i, idx + len(placeholder) - cum
                break
            cum += r_len
        runs[s_run].text = runs[s_run].text[:s_idx] + replacement
        for j in range(s_run + 1, e_run):
            runs[j].text = ''
        if e_run != s_run:
            runs[e_run].text = runs[e_run].text[e_idx:]
        full_text = ''.join(r.text for r in runs)
        idx = full_text.find(placeholder)
# ---------------------------------------

def build(template_path: Path, json_path: Path, out_stub: Path):
    # 1. placeholder map
    content = json.loads(json_path.read_text(encoding="utf-8"))

    # 2. fill template
    doc = Document(template_path)
    paragraphs = get_all_paragraphs(doc)
    for ph, val in content.items():
        for p in paragraphs:
            replace_text_keeping_style(p, ph, val)

    # 3. save DOCX
    docx_path = out_stub.with_suffix(".docx")
    doc.save(docx_path)

    # 4. try to create PDF
    pdf_path = out_stub.with_suffix(".pdf")
    pdf_success = False

    # -- first choice: docx2pdf --
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        pdf_success = pdf_path.exists()
    except ModuleNotFoundError:
        pass  # docx2pdf not installed
    except Exception:
        pass  # docx2pdf failed, fall back below

    # -- second choice: LibreOffice CLI (mac path) --
    if not pdf_success:
        soffice = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        if soffice.exists():
            try:
                subprocess.check_call([
                    str(soffice), "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(pdf_path.parent),
                    str(docx_path)
                ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                pdf_success = pdf_path.exists()
            except Exception:
                pdf_success = False

    if not pdf_success:
        pdf_path = None  # gracefully degrade

    return docx_path, pdf_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True, type=Path)
    ap.add_argument("--json",     required=True, type=Path)
    ap.add_argument("--out",      required=True, type=Path,
                    help="output stub (no extension), e.g. /tmp/myresume")
    args = ap.parse_args()

    try:
        docx_path, pdf_path = build(args.template, args.json, args.out)
        print("✅ DOCX:", docx_path)
        if pdf_path:
            print("✅ PDF :", pdf_path)
        else:
            print("⚠️  PDF conversion skipped (docx2pdf/LibreOffice not found)")
    except Exception as e:
        print(f"❌ Build failed: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
